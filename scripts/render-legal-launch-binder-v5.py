"""Legal Launch Binder — Version 5.

Rebuild of the Attorney Review Packet as a professional legal
diligence binder, suitable for outside counsel, Rocket Lawyer,
investors, and due-diligence review.

Design principles (per user directive):
  - Preserve all policy language exactly; source of truth is the
    Attorney Review Packet HTML extract (/tmp/attorney_packet.json).
  - Use Word Styles for every element (Heading 1/2/3, Body, Quote,
    AttorneyNote, Warning, Appendix, Table, Caption, Code).
  - Auto Table of Contents (Word TOC field — press F9 in Word to
    populate on first open).
  - Automatic page numbering "Page X of Y" via PAGE / NUMPAGES fields.
  - Running header + footer on every page after the cover.
  - Standardized per-policy structure: Divider · Overview · Body ·
    Revision History · Attorney Review Notes.

Env vars:
  IN_FILE  — path to /tmp/attorney_packet.json (default)
  OUT_FILE — /app/frontend/public/downloads/legal-launch-binder-v5-<DATE>.docx
"""
from __future__ import annotations

import json
import os
import re
from datetime import date
from typing import Optional

from bs4 import BeautifulSoup, Tag
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

IN_FILE = os.environ.get("IN_FILE", "/tmp/attorney_packet.json")
OUT_FILE = os.environ.get(
    "OUT_FILE",
    f"/app/frontend/public/downloads/legal-launch-binder-v5-{date.today().isoformat()}.docx",
)

INSPECTOR_ATTR = re.compile(
    r'\s(?:x-source-[a-z-]+|x-file-[a-z-]+|x-line-number|x-column|x-array-var|x-array-index)="[^"]*"'
)

BINDER_VERSION = "5.1"
BINDER_DATE = "June 30, 2026"
CONFIDENTIAL = "CONFIDENTIAL — Attorney Work Product"

# Icons for attorney note categories (Unicode, render in every Word install)
CATEGORY_ICON = {
    "Critical":                 "●",  # red filled circle
    "Counsel Decision Required":"◆",  # purple diamond
    "Recommended":              "▶",  # blue right-triangle
    "Informational":            "○",  # grey open circle
    "Implemented":              "✓",  # green check
}

# Status badges for Launch Readiness Dashboard
STATUS_BADGE = {
    "Critical":         "● Critical",
    "Counsel Required": "◆ Counsel Review",
    "Pending Review":   "◐ Pending",
    "Ready":            "✓ Ready",
}

# Binder version history
BINDER_HISTORY = [
    ("1.0", "Initial draft",      "Initial policy set"),
    ("2.0", "Rocket Lawyer R1",   "Added arbitration, class-action waiver, small-claims carve-out"),
    ("3.0", "AI Policy",          "Added Creator-Owned AI Policy (Ops AI vs. Model Training)"),
    ("4.0", "Counsel review",     "Added Appendix A per policy, non-waivable rights carve-outs"),
    ("5.0", "Binder rebuild",     "New architecture, dashboard, risk matrix, standardized structure"),
    ("5.1", "Polish pass",        "Rich dividers, pre-populated TOC, statistics, formal sign-off"),
]

# ---------------------------------------------------------------------------
# Bookmark + hyperlink helpers
# ---------------------------------------------------------------------------

_BM_ID_COUNTER = [1000]


def _next_bm_id() -> int:
    _BM_ID_COUNTER[0] += 1
    return _BM_ID_COUNTER[0]


def _slug(text: str) -> str:
    s = re.sub(r"[^A-Za-z0-9]+", "_", text).strip("_")
    return "bm_" + s[:40]


def add_bookmark(paragraph, name: str) -> str:
    """Wrap the paragraph's start position with a Word bookmark."""
    bm_id = str(_next_bm_id())
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), bm_id)
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), bm_id)
    p_elem = paragraph._p
    p_elem.insert(0, start)
    p_elem.append(end)
    return name


def add_internal_hyperlink(paragraph, anchor: str, text: str,
                           *, size: int = 11, bold: bool = False,
                           color_hex: str = "1F6FEB") -> None:
    """Add a clickable in-document hyperlink to a bookmark."""
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    hyperlink.set(qn("w:history"), "1")

    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), color_hex)
    rPr.append(color)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(size * 2))
    rPr.append(sz)
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Calibri")
    rFonts.set(qn("w:hAnsi"), "Calibri")
    rPr.append(rFonts)
    if bold:
        b = OxmlElement("w:b")
        rPr.append(b)
    r.append(rPr)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    r.append(t)
    hyperlink.append(r)
    paragraph._p.append(hyperlink)

# ---------------------------------------------------------------------------
# Styles
# ---------------------------------------------------------------------------

def _shade(paragraph, color_hex: str) -> None:
    """Apply background shading to a paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    pPr.append(shd)


def _left_border(paragraph, color_hex: str, sz: str = "24") -> None:
    """Apply a thick left border to a paragraph — used for callouts."""
    pPr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), sz)
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), color_hex)
    pbdr.append(left)
    pPr.append(pbdr)


def build_styles(doc: Document) -> None:
    styles = doc.styles

    # Normal / Body default — Calibri for corporate readability
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    # Standard heading styles — used by Nav Pane + auto-TOC
    for level, size in [(1, 22), (2, 16), (3, 13)]:
        h = styles[f"Heading {level}"]
        h.font.name = "Calibri"
        h.font.size = Pt(size)
        h.font.bold = True
        h.font.color.rgb = RGBColor(0x1F, 0x2A, 0x44)
        h.paragraph_format.space_before = Pt(14 if level == 1 else 10)
        h.paragraph_format.space_after = Pt(4)
        h.paragraph_format.keep_with_next = True

    # Custom binder styles
    def add(name: str, base: str, **kwargs):
        try:
            s = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        except (ValueError, KeyError):
            s = styles[name]
        s.base_style = styles[base]
        f = s.font
        pf = s.paragraph_format
        for k, v in kwargs.items():
            if k == "size":
                f.size = Pt(v)
            elif k == "bold":
                f.bold = v
            elif k == "italic":
                f.italic = v
            elif k == "color":
                f.color.rgb = v
            elif k == "space_after":
                pf.space_after = Pt(v)
            elif k == "space_before":
                pf.space_before = Pt(v)
            elif k == "left_indent":
                pf.left_indent = Cm(v)
            elif k == "font":
                f.name = v
        return s

    add("Body", "Normal", size=11)
    add("Quote", "Normal", size=10, italic=True, left_indent=0.75,
        color=RGBColor(0x33, 0x66, 0x99))
    add("AttorneyNote", "Normal", size=10, left_indent=0.5,
        color=RGBColor(0x33, 0x33, 0x33))
    add("Warning", "Normal", size=10, bold=True, left_indent=0.5,
        color=RGBColor(0xB4, 0x1D, 0x1D))
    add("Appendix", "Normal", size=10, left_indent=0.4,
        color=RGBColor(0x33, 0x33, 0x33))
    add("TableText", "Normal", size=10)
    add("Caption", "Normal", size=9, italic=True,
        color=RGBColor(0x55, 0x55, 0x55), space_after=8)
    add("Code", "Normal", size=10, font="Consolas",
        color=RGBColor(0x33, 0x33, 0x33))
    add("CoverTitle", "Normal", size=32, bold=True,
        color=RGBColor(0x1F, 0x2A, 0x44), space_after=6)
    add("CoverSubtitle", "Normal", size=16,
        color=RGBColor(0x1F, 0x2A, 0x44), space_after=4)
    add("CoverMeta", "Normal", size=11,
        color=RGBColor(0x55, 0x55, 0x55), space_after=2)
    add("DividerLabel", "Normal", size=11, bold=True,
        color=RGBColor(0x1F, 0x2A, 0x44), space_after=2)
    add("DividerTitle", "Normal", size=28, bold=True,
        color=RGBColor(0x1F, 0x2A, 0x44), space_after=12)
    add("DividerMeta", "Normal", size=11,
        color=RGBColor(0x33, 0x33, 0x33), space_after=2)


# ---------------------------------------------------------------------------
# Field / TOC helpers
# ---------------------------------------------------------------------------

def _add_field(paragraph, instr_text: str, *, dirty: bool = False) -> None:
    """Insert a raw Word field (for TOC, PAGE, NUMPAGES, etc.).

    If ``dirty=True``, marks the field as needing recalculation so Word
    prompts to update it on open.
    """
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    if dirty:
        fld_begin.set(qn("w:dirty"), "true")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instr_text
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    # placeholder shown until user refreshes fields in Word (F9)
    placeholder = OxmlElement("w:t")
    placeholder.text = "[Update TOC in Word: right-click → Update Field]"
    run._r.append(placeholder)
    run._r.append(fld_end)


def _enable_auto_update_fields(doc: Document) -> None:
    """Add ``<w:updateFields w:val="true"/>`` to settings.xml so Word
    prompts to refresh TOC / PAGE / NUMPAGES on document open."""
    settings = doc.settings.element
    tag = qn("w:updateFields")
    existing = settings.find(tag)
    if existing is not None:
        existing.set(qn("w:val"), "true")
        return
    el = OxmlElement("w:updateFields")
    el.set(qn("w:val"), "true")
    settings.append(el)


def _add_page_break(doc: Document) -> None:
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def _add_section_break_next_page(doc: Document):
    """Insert a section break so header/footer settings can differ."""
    from docx.enum.section import WD_SECTION
    return doc.add_section(WD_SECTION.NEW_PAGE)


# ---------------------------------------------------------------------------
# Header / footer
# ---------------------------------------------------------------------------

def _apply_header_footer(section, *, include_page_x_of_y: bool = True,
                        show_running_header: bool = True) -> None:
    # Header
    hdr = section.header
    hdr.is_linked_to_previous = False
    hdr_p = hdr.paragraphs[0]
    hdr_p.clear() if hasattr(hdr_p, "clear") else None
    for r in list(hdr_p.runs):
        r.text = ""
    if show_running_header:
        run = hdr_p.add_run(
            "CRAFTERS MARKET   ·   Trust & Policy Center   ·   Legal Launch Binder"
        )
        run.font.name = "Calibri"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    hdr_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # bottom border on header
    pPr = hdr_p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), "AAAAAA")
    pbdr.append(bottom)
    pPr.append(pbdr)

    # Footer
    ftr = section.footer
    ftr.is_linked_to_previous = False
    ftr_p = ftr.paragraphs[0]
    for r in list(ftr_p.runs):
        r.text = ""
    left = ftr_p.add_run("Confidential  ·  Prepared for Rocket Lawyer     ")
    left.font.name = "Calibri"
    left.font.size = Pt(9)
    left.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    if include_page_x_of_y:
        _add_field(ftr_p, " PAGE ")
        mid = ftr_p.add_run(" of ")
        mid.font.name = "Calibri"
        mid.font.size = Pt(9)
        mid.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        _add_field(ftr_p, " NUMPAGES ")

    ftr_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT


# ---------------------------------------------------------------------------
# Helpers to add styled content
# ---------------------------------------------------------------------------

def p(doc, text: str, style: str = "Body") -> object:
    para = doc.add_paragraph(style=style)
    para.add_run(text)
    return para


def h(doc, text: str, level: int, *, bookmark: Optional[str] = None) -> object:
    para = doc.add_heading(text, level=level)
    if bookmark or level == 1:
        add_bookmark(para, bookmark or _slug(text))
    return para


def bullet(doc, text: str) -> object:
    return doc.add_paragraph(text, style="List Bullet")


def numbered(doc, text: str) -> object:
    return doc.add_paragraph(text, style="List Number")


def callout(doc, text: str, category: str = "Recommended") -> None:
    """Attorney callout box — colored left border + shaded background + icon."""
    palette = {
        "Critical":                 ("B41D1D", "FDECEC", "CRITICAL"),
        "Counsel Decision Required":("8A2BE2", "F3EAFB", "COUNSEL DECISION REQUIRED"),
        "Recommended":              ("1F6FEB", "E8F0FE", "RECOMMENDED"),
        "Informational":            ("6C757D", "F1F3F5", "INFORMATIONAL"),
        "Implemented":              ("1E8E3E", "E7F5EC", "IMPLEMENTED"),
    }
    accent, fill, label = palette.get(category, palette["Recommended"])
    icon = CATEGORY_ICON.get(category, "○")

    label_p = doc.add_paragraph(style="AttorneyNote")
    label_run = label_p.add_run(f"{icon}  {label}")
    label_run.bold = True
    label_run.font.size = Pt(9)
    label_run.font.color.rgb = RGBColor.from_string(accent)
    _left_border(label_p, accent)
    _shade(label_p, fill)

    body_p = doc.add_paragraph(style="AttorneyNote")
    body_p.add_run(text).font.size = Pt(10)
    _left_border(body_p, accent)
    _shade(body_p, fill)
    body_p.paragraph_format.space_after = Pt(10)


# ---------------------------------------------------------------------------
# Extraction from source HTML
# ---------------------------------------------------------------------------

def _text(node) -> str:
    if node is None:
        return ""
    return re.sub(r"\s+", " ", node.get_text(" ", strip=True)).strip()


def _has_class(tag: Tag, cls: str) -> bool:
    if not isinstance(tag, Tag):
        return False
    classes = tag.get("class", [])
    return cls in (classes if isinstance(classes, list) else classes.split())


def load_source() -> BeautifulSoup:
    with open(IN_FILE) as f:
        data = json.load(f)
    html = INSPECTOR_ATTR.sub("", data["html"])
    return BeautifulSoup(html, "lxml")


def extract_policies(soup: BeautifulSoup) -> list[dict]:
    out = []
    for i, pol in enumerate(soup.select(".pkt-policy"), start=1):
        title = _text(pol.select_one(".pkt-h1-pol") or pol.select_one("h1"))
        meta = _text(pol.select_one(".pkt-meta-inline"))
        description = _text(pol.select_one(".pkt-description"))
        intro = _text(pol.select_one(".pkt-intro"))

        version = ""
        effective = ""
        last_updated = ""
        # Meta format like: "Version: 2.6 · Effective: 2026-06-30 · Last Updated: 2026-06-30"
        m = re.search(r"Version[:\s]+([0-9.]+)", meta)
        if m: version = m.group(1)
        m = re.search(r"Effective[:\s]+([\w\s\-,]+?)(?:·|$)", meta)
        if m: effective = m.group(1).strip()
        m = re.search(r"Last [Uu]pdated[:\s]+([\w\s\-,]+?)(?:·|$)", meta)
        if m: last_updated = m.group(1).strip()

        # Sections (numbered body blocks)
        blocks = []
        for bb in pol.select(".pkt-body-block"):
            heading = _text(bb.select_one(".pkt-h3, .pkt-h4, h3, h4"))
            body_parts = []
            for el in bb.children:
                if not isinstance(el, Tag): continue
                if _has_class(el, "pkt-h3") or _has_class(el, "pkt-h4") or el.name in ("h3", "h4"):
                    continue
                if el.name == "p":
                    txt = _text(el)
                    if txt: body_parts.append(("p", txt))
                elif el.name == "ul":
                    for li in el.find_all("li", recursive=False):
                        body_parts.append(("bullet", _text(li)))
                elif el.name == "ol":
                    for li in el.find_all("li", recursive=False):
                        body_parts.append(("numbered", _text(li)))
                elif el.name == "dl":
                    for row in el.find_all(["dt", "dd"], recursive=False):
                        body_parts.append(("dt" if row.name == "dt" else "dd", _text(row)))
            blocks.append({"heading": heading, "parts": body_parts})

        # Final callout (info banner)
        callout_text = _text(pol.select_one(".pkt-callout"))

        # Revision history
        rev = []
        rev_block = pol.select_one(".pkt-revision")
        if rev_block:
            for li in rev_block.find_all("li"):
                rev.append(_text(li))

        # Related policies
        related = []
        rel_block = pol.select_one(".pkt-related")
        if rel_block:
            for li in rel_block.find_all("li"):
                related.append(_text(li))

        # Attorney review notes
        attorney_notes = []
        att_block = pol.select_one(".pkt-attorney")
        if att_block:
            for li in att_block.find_all("li"):
                attorney_notes.append(_text(li))

        out.append({
            "index": i,
            "title": title,
            "version": version,
            "effective": effective,
            "last_updated": last_updated,
            "description": description,
            "intro": intro,
            "blocks": blocks,
            "callout": callout_text,
            "revision": rev,
            "related": related,
            "attorney_notes": attorney_notes,
        })
    return out


def categorize_note(note: str) -> str:
    """Best-effort categorization of an attorney note."""
    up = note.upper()
    if up.startswith("IMPLEMENTED"):
        return "Implemented"
    if "CRITICAL" in up or "MUST" in up or "REQUIRED BEFORE LAUNCH" in up:
        return "Critical"
    if "COUNSEL" in up and ("DECIDE" in up or "CONFIRM" in up or "APPROVE" in up):
        return "Counsel Decision Required"
    if "CONFIRM" in up or "REVIEW" in up or "VERIFY" in up:
        return "Recommended"
    return "Informational"


# ---------------------------------------------------------------------------
# Binder sections
# ---------------------------------------------------------------------------

def build_cover(doc: Document) -> None:
    for _ in range(3):
        doc.add_paragraph()

    # Top accent band
    band_top = doc.add_paragraph()
    _shade(band_top, "1F2A44")
    band_top.paragraph_format.space_after = Pt(2)
    band_top.add_run(" ").font.size = Pt(2)

    # Brand name
    brand = doc.add_paragraph(style="CoverMeta")
    brand.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = brand.add_run("CRAFTERS  MARKET")
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor(0x1F, 0x2A, 0x44)

    # Big title
    title = doc.add_paragraph(style="CoverTitle")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Trust & Policy Center")

    sub = doc.add_paragraph(style="CoverSubtitle")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run("Legal Launch Binder")

    # Version badge
    for _ in range(2): doc.add_paragraph()
    ver = doc.add_paragraph()
    ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _shade(ver, "E8F0FE")
    ver.paragraph_format.space_after = Pt(0)
    r = ver.add_run(f"  VERSION  {BINDER_VERSION}  ")
    r.bold = True; r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0x1F, 0x6F, 0xEB)

    # Body meta lines
    doc.add_paragraph()
    for line in [
        "Prepared for Rocket Lawyer",
        "Prepared by Crafters Market Operations",
        BINDER_DATE,
    ]:
        para = doc.add_paragraph(style="CoverMeta")
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.add_run(line)

    for _ in range(6): doc.add_paragraph()

    # Bottom accent band
    band_bot = doc.add_paragraph()
    _shade(band_bot, "1F2A44")
    band_bot.paragraph_format.space_after = Pt(2)
    band_bot.add_run(" ").font.size = Pt(2)

    # Classification badge
    conf = doc.add_paragraph()
    conf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _shade(conf, "FDECEC")
    r = conf.add_run(f"  {CONFIDENTIAL}  ")
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0xB4, 0x1D, 0x1D)


def build_binder_statistics(doc: Document, policies: list[dict]) -> None:
    h(doc, "Binder Statistics", 1, bookmark="bm_Binder_Statistics")
    p(doc, "Executive snapshot of the diligence surface area covered by this binder.", style="Caption")

    total_notes = sum(len(p["attorney_notes"]) for p in policies)
    by_cat = {"Critical": 0, "Counsel Decision Required": 0,
              "Recommended": 0, "Informational": 0, "Implemented": 0}
    for p_ in policies:
        for n in p_["attorney_notes"]:
            by_cat[categorize_note(n)] += 1

    rows = [
        ("Policies",                             str(len(policies))),
        ("Attorney Review Notes (total)",        str(total_notes)),
        ("Critical Items",                       f"{CATEGORY_ICON['Critical']}  {by_cat['Critical']}"),
        ("Counsel Decisions Required",           f"{CATEGORY_ICON['Counsel Decision Required']}  {by_cat['Counsel Decision Required']}"),
        ("Recommended Reviews",                  f"{CATEGORY_ICON['Recommended']}  {by_cat['Recommended']}"),
        ("Informational Notes",                  f"{CATEGORY_ICON['Informational']}  {by_cat['Informational']}"),
        ("Implemented Recommendations",          f"{CATEGORY_ICON['Implemented']}  {by_cat['Implemented']}"),
    ]
    t = doc.add_table(rows=len(rows), cols=2)
    t.style = "Light Grid Accent 1"
    for i, (k, v) in enumerate(rows):
        row = t.rows[i].cells
        row[0].text = ""
        r = row[0].paragraphs[0].add_run(k)
        r.bold = True; r.font.size = Pt(11)
        row[1].text = ""
        r = row[1].paragraphs[0].add_run(v)
        r.font.size = Pt(11)


def build_binder_version_history(doc: Document) -> None:
    h(doc, "Binder Version History", 1, bookmark="bm_Binder_Version_History")
    p(doc, "Evolution of the Trust & Policy Center from initial draft to the current diligence-grade binder.", style="Caption")
    t = doc.add_table(rows=1 + len(BINDER_HISTORY), cols=3)
    t.style = "Light Grid Accent 1"
    for i, hdr in enumerate(["Version", "Milestone", "Summary"]):
        cell = t.rows[0].cells[i]
        cell.text = ""
        r = cell.paragraphs[0].add_run(hdr)
        r.bold = True; r.font.size = Pt(10)
    for i, (ver, milestone, summary) in enumerate(BINDER_HISTORY, start=1):
        row = t.rows[i].cells
        row[0].text = ""
        r = row[0].paragraphs[0].add_run(f"v{ver}")
        r.bold = True; r.font.size = Pt(10)
        row[1].text = milestone
        row[2].text = summary
        for c in row[1:]:
            for para in c.paragraphs:
                for r in para.runs:
                    r.font.size = Pt(10)


def build_confidentiality(doc: Document) -> None:
    h(doc, "Confidentiality Notice", 1)
    p(doc, "This document, and the appendices attached hereto, contain "
        "confidential and proprietary information of Crafters Market prepared "
        "for the purpose of legal review. It is intended solely for the use of "
        "outside counsel and authorized personnel of Crafters Market.")
    p(doc, "This binder constitutes attorney work-product and is subject to "
        "the attorney–client privilege. Redistribution to any third party "
        "without the express written consent of Crafters Market is prohibited.")
    p(doc, "If you have received this document in error, please notify "
        "policy@craftersmarket.org immediately and destroy all copies.")


def build_executive_summary(doc: Document) -> None:
    h(doc, "Executive Summary", 1)

    h(doc, "Marketplace Snapshot", 2)
    p(doc, "Crafters Market is a U.S.-based, curated online marketplace "
        "connecting independent Makers with Buyers of handmade, handcrafted, "
        "and designer goods. Makers list, price, describe, and fulfill their "
        "own Orders; Crafters Market operates the platform, processes payments "
        "via Stripe Connect, and provides marketplace-level trust, safety, and "
        "customer-support layers.")

    h(doc, "Business Model", 2)
    bullet(doc, "Two tiers: Free (5% commission + 3% processing) and Crafters Plus (4% commission + 3% processing, $12/month).")
    bullet(doc, "Listing fees of $0.20 after the tier-specific free allowance.")
    bullet(doc, "Off-site advertising fee of 12% on sales attributed to Crafters Market ad campaigns.")
    bullet(doc, "Optional Promoted Listings at $5 per week per Listing.")
    bullet(doc, "Founding Seller Program provides cohort-specific benefits layered on top of the Maker's tier.")

    h(doc, "Marketplace Facilitator Posture", 2)
    p(doc, "In U.S. states that treat online marketplaces as marketplace "
        "facilitators, Crafters Market collects and remits sales tax on qualifying "
        "Orders on the Maker's behalf. In other jurisdictions the Maker remains "
        "responsible for tax collection and remittance.")

    h(doc, "Stripe Connect", 2)
    p(doc, "All payments and payouts flow through Stripe Connect. Makers "
        "onboard to Stripe Connect for KYC, banking, and payout scheduling. "
        "Certain payout holds may be imposed directly by Stripe, payment "
        "networks, financial institutions, or regulatory authorities; Crafters "
        "Market cannot override those holds where it does not control fund "
        "release.")

    h(doc, "Governing Law", 2)
    p(doc, "The Terms of Service and Maker Agreement are governed by the laws "
        "of the State of Washington. Disputes are resolved through mandatory, "
        "individual arbitration administered by the American Arbitration "
        "Association, seated in King County, Washington, and conducted "
        "remotely by default (video conference or written submissions), unless "
        "the arbitrator determines that an in-person hearing is necessary. "
        "Small-claims and injunctive-relief claims are carved out.")

    h(doc, "Creator-Owned AI Policy", 2)
    p(doc, "Operational AI (search, recommendations, moderation, listing "
        "optimization, translation, off-site ad generation) is permitted under "
        "the User Content License. AI Model Training on Maker Content is not "
        "authorized under that license and requires separate opt-in consent. "
        "Operational AI does not authorize the Platform or any third-party "
        "advertising provider to train commercial foundation models on Maker "
        "Content.")

    h(doc, "Review Objectives", 2)
    numbered(doc, "Confirm enforceability of the arbitration clause and class-action waiver.")
    numbered(doc, "Confirm the liability cap ($100 / 12-month commission floor) is appropriate for a Version-1.0 marketplace.")
    numbered(doc, "Confirm indemnification carve-outs (gross negligence, willful misconduct) satisfy prevailing practice.")
    numbered(doc, "Confirm the AI Operational vs. Model Training distinction is enforceable and adequately disclosed.")
    numbered(doc, "Confirm marketplace-facilitator posture is aligned with active state regimes.")
    numbered(doc, "Sign off on the Fee & Pricing Policy as the single source of truth for commercial terms.")

    h(doc, "Requested Deliverables", 2)
    bullet(doc, "Line-by-line redlines against each of the 14 policies.")
    bullet(doc, "Overall sign-off, or itemized blockers with proposed language.")
    bullet(doc, "Confirmation of Attorney Review Notes (Appendix A per policy, plus the Master Appendix A at the end).")
    bullet(doc, "Completion of the Counsel Workbook at the end of the binder.")


def build_document_control(doc: Document) -> None:
    h(doc, "Document Control", 1)
    rows = [
        ("Binder Version", f"{BINDER_VERSION}"),
        ("Owner", "Crafters Market Operations · policy@craftersmarket.org"),
        ("Classification", "Confidential — Attorney Work Product"),
        ("Review Cycle", "Standing quarterly review + ad-hoc for material changes"),
        ("Effective Date", "Set at production launch (see individual policies)"),
        ("Revision", f"Binder v{BINDER_VERSION} — {BINDER_DATE}"),
        ("Approval Status", "Awaiting counsel sign-off"),
        ("Source of Truth", "Crafters Market Trust & Policy Center (this binder is a snapshot)"),
    ]
    t = doc.add_table(rows=len(rows), cols=2)
    t.style = "Light Grid Accent 1"
    for i, (k, v) in enumerate(rows):
        row = t.rows[i].cells
        row[0].text = ""
        row[0].paragraphs[0].add_run(k).bold = True
        row[0].paragraphs[0].runs[0].font.size = Pt(10)
        row[1].text = ""
        r = row[1].paragraphs[0].add_run(v)
        r.font.size = Pt(10)


def build_marketplace_overview(doc: Document) -> None:
    h(doc, "Marketplace Overview", 1)
    p(doc, "Crafters Market is a curated U.S. online marketplace for handmade "
        "and handcrafted goods. Makers are vetted, approved, and admitted to "
        "the Platform on a rolling basis. Buyers browse, purchase, and receive "
        "goods shipped directly by the Maker. Crafters Market operates the "
        "search, discovery, checkout, payment, dispute-resolution, and "
        "customer-support layers on top of the Maker's own shop.")
    p(doc, "The Platform launches in Version 1.0 as a U.S.-focused service. "
        "International expansion (EEA, UK, and jurisdictions with additional "
        "transfer safeguards) will be evaluated post-launch, with the "
        "corresponding privacy, consumer-protection, and tax safeguards "
        "implemented before accepting registrations from those jurisdictions.")


def build_legal_review_scope(doc: Document) -> None:
    h(doc, "Legal Review Scope", 1)
    p(doc, "This binder consolidates the fourteen (14) policies that together "
        "constitute the Crafters Market Trust & Policy Center. Each policy is "
        "presented with a standardized front-matter (Purpose, Scope, Applies "
        "To, Dependencies, Attorney Focus), the full policy body, its "
        "revision history, and its Attorney Review Notes (Appendix A). A "
        "consolidated Master Appendix A grouped by priority appears at the "
        "end of the binder.")
    bullet(doc, "Core Policies: Terms of Service, Maker Agreement, Privacy Policy, Buyer Protection Policy, Returns & Refunds Policy, Shipping & Logistics Policy, Prohibited Items Policy, Community Guidelines, Intellectual Property & DMCA Policy, Cookies Policy.")
    bullet(doc, "Operational Policies: Fee & Pricing Policy, Marketplace Promise, Accessibility Statement.")
    bullet(doc, "Cross-cutting: Creator-Owned AI Policy (integrated into Terms §6a, Maker Agreement §10a, Privacy §11).")


def build_launch_readiness_dashboard(doc: Document, policies: list[dict]) -> None:
    h(doc, "Launch Readiness Dashboard", 1)
    p(doc, "Snapshot of every policy and its readiness posture at binder-issue date.", style="Caption")
    t = doc.add_table(rows=1 + len(policies), cols=5)
    t.style = "Light Grid Accent 1"
    headers = ["#", "Policy", "Version", "Attorney Focus", "Readiness"]
    for i, hdr in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        r = cell.paragraphs[0].add_run(hdr)
        r.bold = True; r.font.size = Pt(10)
    for i, pol in enumerate(policies, start=1):
        row = t.rows[i].cells
        row[0].text = f"{i:02d}"
        row[1].text = pol["title"]
        row[2].text = pol["version"] or "—"
        first_note = pol["attorney_notes"][0] if pol["attorney_notes"] else ""
        row[3].text = (first_note[:120] + "…") if len(first_note) > 120 else (first_note or "—")
        # Readiness: if any Critical -> Critical, if any note remaining -> Pending, else Ready
        cats = [categorize_note(n) for n in pol["attorney_notes"]]
        if "Critical" in cats: row[4].text = STATUS_BADGE["Critical"]
        elif "Counsel Decision Required" in cats: row[4].text = STATUS_BADGE["Counsel Required"]
        elif cats: row[4].text = STATUS_BADGE["Pending Review"]
        else: row[4].text = STATUS_BADGE["Ready"]
        for c in row:
            for para in c.paragraphs:
                for r in para.runs:
                    r.font.size = Pt(9)


def build_risk_matrix(doc: Document, policies: list[dict]) -> None:
    h(doc, "Risk Matrix", 1)
    p(doc, "Executive matrix — one row per policy, categorized attorney notes.", style="Caption")
    t = doc.add_table(rows=1 + len(policies), cols=5)
    t.style = "Light Grid Accent 1"
    for i, hdr in enumerate(["Policy", "Status", "Attorney Review", "Priority", "Comments"]):
        cell = t.rows[0].cells[i]
        cell.text = ""
        r = cell.paragraphs[0].add_run(hdr)
        r.bold = True; r.font.size = Pt(10)
    for i, pol in enumerate(policies, start=1):
        cats = [categorize_note(n) for n in pol["attorney_notes"]]
        priority = ("Critical" if "Critical" in cats else
                    "High" if "Counsel Decision Required" in cats else
                    "Medium" if "Recommended" in cats else
                    "Low")
        status = "Pending Review" if pol["attorney_notes"] else "Ready"
        review = f"{len(pol['attorney_notes'])} note(s)"
        comments = pol["attorney_notes"][0][:90] + "…" if pol["attorney_notes"] and len(pol["attorney_notes"][0]) > 90 else (pol["attorney_notes"][0] if pol["attorney_notes"] else "—")
        for j, val in enumerate([pol["title"], status, review, priority, comments]):
            cell = t.rows[i].cells[j]
            cell.text = ""
            r = cell.paragraphs[0].add_run(val)
            r.font.size = Pt(9)


def build_counsel_deliverables(doc: Document) -> None:
    h(doc, "Counsel Deliverables", 1)
    numbered(doc, "Redlines against each of the 14 policies in this binder (attach as tracked-changes DOCX or PDF markup).")
    numbered(doc, "Confirmation on every open Attorney Review Note in Appendix A per policy.")
    numbered(doc, "Overall Launch Recommendation (Approve / Approve with Redlines / Do Not Launch).")
    numbered(doc, "Signed Attorney Sign-off block at the end of this binder.")


def build_open_legal_questions(doc: Document, policies: list[dict]) -> None:
    h(doc, "Open Legal Questions", 1)
    p(doc, "The following questions remain open at binder-issue date. Full "
        "context for each question is preserved in the policy-level Appendix A "
        "and consolidated in the Master Appendix A at the end of this binder.",
        style="Caption")
    for pol in policies:
        if not pol["attorney_notes"]:
            continue
        h(doc, pol["title"], 3)
        for n in pol["attorney_notes"]:
            cat = categorize_note(n)
            callout(doc, n, cat)


def build_cross_policy_dependency_map(doc: Document) -> None:
    h(doc, "Cross-Policy Dependency Map", 1)
    p(doc, "Canonical order of precedence, applied consistently across every "
        "policy in this binder.", style="Caption")
    steps = [
        ("Level 1", "Applicable Law",
         "Non-waivable consumer-protection rights always govern."),
        ("Level 2", "Terms of Service",
         "Foundational contract between every User and the Platform."),
        ("Level 3", "Maker Agreement (seller-specific issues only)",
         "For issues relating to Maker activity (listings, payouts, seller IP, exclusivity, taxes)."),
        ("Level 4", "Marketplace Policies",
         "Buyer Protection · Returns · Shipping · Privacy · Cookies · Prohibited Items · Community Guidelines · IP/DMCA · Fee & Pricing · Accessibility."),
        ("Level 5", "Maker Shop Policies",
         "A Maker's own published Shop Policies. Must not conflict with anything above."),
        ("Level 6", "Order-Specific Agreements",
         "Terms agreed at checkout or in messaging for a specific Order. Bind only that Order."),
    ]
    t = doc.add_table(rows=len(steps), cols=3)
    t.style = "Light Grid Accent 1"
    for i, (lvl, label, note) in enumerate(steps):
        row = t.rows[i].cells
        row[0].text = ""
        r = row[0].paragraphs[0].add_run(lvl)
        r.bold = True; r.font.size = Pt(10)
        row[1].text = ""
        r2 = row[1].paragraphs[0].add_run(label)
        r2.bold = True; r2.font.size = Pt(10)
        row[2].text = ""
        r3 = row[2].paragraphs[0].add_run(note)
        r3.font.size = Pt(9)


def build_toc(doc: Document, policies: list[dict]) -> None:
    heading = h(doc, "Table of Contents", 1)
    add_bookmark(heading, "bm_Table_of_Contents")
    p(doc, "The Table of Contents below is generated from the binder's "
      "Heading 1/2/3 styles. It is fully populated in the distribution PDF; "
      "in the editable DOCX, click anywhere inside the TOC and press F9 "
      "(or right-click → 'Update Field' → 'Update entire table') after any "
      "edits to refresh page numbers. A supplementary Hyperlinked Navigation "
      "Index follows for quick cross-referencing to top-level sections.",
      style="Caption")

    # --- Word auto TOC field (Heading 1-3) --------------------------------
    toc_para = doc.add_paragraph()
    toc_para.paragraph_format.space_after = Pt(6)
    _add_field(
        toc_para,
        ' TOC \\o "1-3" \\h \\z \\u ',
        dirty=True,
    )

    _add_page_break(doc)

    # --- Hyperlinked navigation index (pre-populated fallback) ------------
    subhead = doc.add_paragraph(style="Heading 2")
    subhead.add_run("Hyperlinked Navigation Index").bold = True

    p(doc, "Every entry below is a clickable internal hyperlink to the "
      "corresponding bookmark. Use these if the auto-TOC has not yet been "
      "refreshed in Word.", style="Caption")

    entries = [
        ("Confidentiality Notice",         "bm_Confidentiality_Notice"),
        ("Executive Summary",              "bm_Executive_Summary"),
        ("Binder Statistics",              "bm_Binder_Statistics"),
        ("Document Control",               "bm_Document_Control"),
        ("Binder Version History",         "bm_Binder_Version_History"),
        ("Marketplace Overview",           "bm_Marketplace_Overview"),
        ("Legal Review Scope",             "bm_Legal_Review_Scope"),
        ("Launch Readiness Dashboard",     "bm_Launch_Readiness_Dashboard"),
        ("Risk Matrix",                    "bm_Risk_Matrix"),
        ("Counsel Deliverables",           "bm_Counsel_Deliverables"),
        ("Open Legal Questions",           "bm_Open_Legal_Questions"),
        ("Cross-Policy Dependency Map",    "bm_Cross_Policy_Dependency_Map"),
    ]
    for label, anchor in entries:
        line = doc.add_paragraph()
        line.paragraph_format.space_after = Pt(3)
        add_internal_hyperlink(line, anchor, label, size=11, bold=True)

    # Blank line then policies header
    doc.add_paragraph()
    ph = doc.add_paragraph()
    ph.add_run("Policies").bold = True

    for pol in policies:
        anchor = _slug(f"policy_{pol['index']:02d}_{pol['title']}")
        label = f"Policy {pol['index']:02d} — {pol['title']}"
        line = doc.add_paragraph()
        line.paragraph_format.left_indent = Cm(0.4)
        line.paragraph_format.space_after = Pt(2)
        add_internal_hyperlink(line, anchor, label, size=11)
        version = pol["version"] or "—"
        r = line.add_run(f"    ·  v{version}")
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    # Trailing
    doc.add_paragraph()
    ph2 = doc.add_paragraph()
    ph2.add_run("Appendices & Workbook").bold = True
    for label, anchor in [
        ("Master Appendix A — Consolidated Attorney Review Notes", "bm_Master_Appendix_A"),
        ("Master Revision Log",                                    "bm_Master_Revision_Log"),
        ("Counsel Workbook",                                       "bm_Counsel_Workbook"),
        ("Launch Recommendation",                                  "bm_Launch_Recommendation"),
        ("Attorney Sign-off",                                      "bm_Attorney_Sign_off"),
    ]:
        line = doc.add_paragraph()
        line.paragraph_format.space_after = Pt(3)
        add_internal_hyperlink(line, anchor, label, size=11, bold=True)


# ---------------------------------------------------------------------------
# Policy sections
# ---------------------------------------------------------------------------

def _derive_applies_to(pol: dict) -> str:
    t = pol["title"].lower()
    if "maker" in t or "prohibited" in t or "community" in t:
        return "Makers (sellers) admitted to Crafters Market."
    if "buyer" in t or "shipping" in t or "return" in t:
        return "Buyers who purchase through Crafters Market."
    if "cookie" in t:
        return "Any visitor to the Platform's web surfaces."
    if "accessibility" in t or "marketplace promise" in t or "creator-owned ai" in t:
        return "All Users of the Platform (Buyers, Makers, and visitors)."
    return "All Users of the Platform (Buyers, Makers, and visitors)."


def _derive_attorney_focus(pol: dict) -> str:
    if not pol["attorney_notes"]:
        return "No outstanding attorney focus items at binder-issue date."
    return pol["attorney_notes"][0]


def _priority_label(pol: dict) -> str:
    cats = [categorize_note(n) for n in pol["attorney_notes"]]
    if "Critical" in cats: return "Critical"
    if "Counsel Decision Required" in cats: return "High"
    if "Recommended" in cats: return "Medium"
    if pol["attorney_notes"]: return "Low"
    return "None"


def build_policy_divider(doc: Document, pol: dict) -> None:
    # Full-page divider with rich pre-read metadata
    anchor_name = _slug(f"policy_{pol['index']:02d}_{pol['title']}")

    for _ in range(3):
        doc.add_paragraph()

    # Accent bar (top)
    bar = doc.add_paragraph()
    _shade(bar, "1F2A44")
    bar.paragraph_format.space_after = Pt(4)
    bar.add_run(" ").font.size = Pt(2)

    label = doc.add_paragraph(style="DividerLabel")
    label.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_bookmark(label, anchor_name)
    r = label.add_run(f"POLICY {pol['index']:02d}")
    r.font.size = Pt(12)

    title = doc.add_paragraph(style="DividerTitle")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run(pol["title"].upper())
    tr.font.size = Pt(26)

    # Accent bar (bottom)
    bar2 = doc.add_paragraph()
    _shade(bar2, "1F2A44")
    bar2.paragraph_format.space_after = Pt(12)
    bar2.add_run(" ").font.size = Pt(2)

    # Rich pre-read table
    rows = [
        ("Purpose",         pol["description"] or pol["intro"][:200] or f"See {pol['title']}."),
        ("Scope",           "As defined in the numbered sections that follow, subject to the canonical Cross-Policy Dependency Map at the front of this binder."),
        ("Applies To",      _derive_applies_to(pol)),
        ("Dependencies",    ", ".join(pol["related"]) if pol["related"] else "None beyond the canonical hierarchy."),
        ("Attorney Focus",  _derive_attorney_focus(pol)[:300]),
        ("Risk Level",      _priority_label(pol)),
        ("Version",         pol["version"] or "—"),
        ("Effective Date",  pol["effective"] or "Set at production launch"),
        ("Last Updated",    pol["last_updated"] or "—"),
    ]
    t = doc.add_table(rows=len(rows), cols=2)
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (k, v) in enumerate(rows):
        cell_k, cell_v = t.rows[i].cells
        cell_k.text = ""
        rk = cell_k.paragraphs[0].add_run(k)
        rk.bold = True; rk.font.size = Pt(10)
        cell_v.text = ""
        rv = cell_v.paragraphs[0].add_run(v)
        rv.font.size = Pt(10)


def build_policy_overview(doc: Document, pol: dict) -> None:
    h(doc, "Policy Overview", 2)
    h(doc, "Purpose", 3)
    p(doc, pol["description"] or pol["intro"] or f"See {pol['title']} body.")
    h(doc, "Scope", 3)
    p(doc, "This policy governs the subject-matter defined in its numbered "
      "sections below, subject to the canonical Cross-Policy Dependency Map "
      "at the front of this binder.")
    h(doc, "Applies To", 3)
    p(doc, _derive_applies_to(pol))
    h(doc, "Dependencies", 3)
    if pol["related"]:
        for r in pol["related"]:
            bullet(doc, r)
    else:
        p(doc, "None beyond the canonical hierarchy.")
    h(doc, "Attorney Focus", 3)
    p(doc, _derive_attorney_focus(pol), style="AttorneyNote")


def build_policy_body(doc: Document, pol: dict) -> None:
    h(doc, "Policy", 2)
    if pol["intro"] and pol["intro"] != pol["description"]:
        p(doc, pol["intro"])

    for block in pol["blocks"]:
        if block["heading"]:
            h(doc, block["heading"], 3)
        for kind, txt in block["parts"]:
            if kind == "p":
                p(doc, txt)
            elif kind == "bullet":
                bullet(doc, txt)
            elif kind == "numbered":
                numbered(doc, txt)
            elif kind == "dt":
                para = doc.add_paragraph()
                r = para.add_run(txt)
                r.bold = True; r.font.size = Pt(10)
            elif kind == "dd":
                p(doc, txt)

    if pol["callout"]:
        callout(doc, pol["callout"], "Informational")


def build_policy_revision_history(doc: Document, pol: dict) -> None:
    if not pol["revision"]:
        return
    h(doc, "Revision History", 2)
    for r in pol["revision"]:
        bullet(doc, r)


def build_policy_appendix_a(doc: Document, pol: dict) -> None:
    if not pol["attorney_notes"]:
        return
    h(doc, "Appendix A — Attorney Review Notes", 2)
    for n in pol["attorney_notes"]:
        callout(doc, n, categorize_note(n))


# ---------------------------------------------------------------------------
# Master Appendix A + Revision Log
# ---------------------------------------------------------------------------

def build_master_appendix_a(doc: Document, policies: list[dict]) -> None:
    h(doc, "Master Appendix A — Consolidated Attorney Review Notes", 1)
    groups: dict[str, list[tuple[str, str]]] = {
        "Critical": [], "Counsel Decision Required": [],
        "Recommended": [], "Informational": [], "Implemented": [],
    }
    for pol in policies:
        for n in pol["attorney_notes"]:
            groups[categorize_note(n)].append((pol["title"], n))

    for cat in ["Critical", "Counsel Decision Required", "Recommended",
                "Informational", "Implemented"]:
        entries = groups[cat]
        if not entries:
            continue
        h(doc, cat, 2)
        for policy_title, note in entries:
            cap = doc.add_paragraph(style="Caption")
            cap.add_run(policy_title).bold = True
            callout(doc, note, cat)


def build_master_revision_log(doc: Document, policies: list[dict]) -> None:
    h(doc, "Master Revision Log", 1)
    t = doc.add_table(rows=1 + len(policies), cols=4)
    t.style = "Light Grid Accent 1"
    for i, hdr in enumerate(["#", "Policy", "Version", "Last Updated"]):
        cell = t.rows[0].cells[i]
        cell.text = ""
        r = cell.paragraphs[0].add_run(hdr)
        r.bold = True; r.font.size = Pt(10)
    for i, pol in enumerate(policies, start=1):
        row = t.rows[i].cells
        row[0].text = f"{i:02d}"
        row[1].text = pol["title"]
        row[2].text = pol["version"] or "—"
        row[3].text = pol["last_updated"] or "—"
        for c in row:
            for para in c.paragraphs:
                for r in para.runs:
                    r.font.size = Pt(9)


def build_counsel_workbook(doc: Document, policies: list[dict]) -> None:
    h(doc, "Counsel Workbook", 1)
    p(doc, "Please complete this workbook and return with your redlines. "
      "It gives Crafters Market Operations a single, structured summary of "
      "your review outcome.")

    h(doc, "Overall Recommendation", 2)
    for opt in [
        "Approve for launch — no changes required.",
        "Approve for launch — with the redlines attached.",
        "Do NOT launch — see Critical Issues below.",
    ]:
        bullet(doc, f"☐  {opt}")

    h(doc, "Critical Issues", 2)
    for _ in range(5):
        p(doc, "☐  __________________________________________________________________________")

    h(doc, "Required Changes", 2)
    for _ in range(5):
        p(doc, "☐  __________________________________________________________________________")

    h(doc, "Suggested Improvements", 2)
    for _ in range(5):
        p(doc, "☐  __________________________________________________________________________")

    h(doc, "Per-Policy Sign-off", 2)
    t = doc.add_table(rows=1 + len(policies), cols=5)
    t.style = "Light Grid Accent 1"
    for i, hdr in enumerate(["Policy", "Version", "Sign-off (Y/N)",
                             "Blockers", "Est. hours"]):
        cell = t.rows[0].cells[i]
        cell.text = ""
        r = cell.paragraphs[0].add_run(hdr)
        r.bold = True; r.font.size = Pt(10)
    for i, pol in enumerate(policies, start=1):
        row = t.rows[i].cells
        row[0].text = pol["title"]
        row[1].text = pol["version"] or "—"
        for j in range(2, 5):
            row[j].text = ""


def build_launch_recommendation(doc: Document) -> None:
    h(doc, "Launch Recommendation", 1)
    p(doc, "Circle one, add the effective date, and sign below.")
    for opt in [
        "APPROVED for launch — no material changes required.",
        "APPROVED with REDLINES — see attached tracked-changes.",
        "HELD — do not launch. See Critical Issues in the Counsel Workbook.",
    ]:
        p(doc, f"☐  {opt}", style="AttorneyNote")


def build_attorney_signoff(doc: Document) -> None:
    h(doc, "Attorney Sign-off", 1)
    p(doc, "By signing below, counsel confirms review of this binder "
      "consistent with the Launch Recommendation above.")

    t = doc.add_table(rows=5, cols=2)
    t.style = "Light Grid Accent 1"
    rows = [
        ("Attorney (print name)", ""),
        ("Firm", ""),
        ("Bar ID / Jurisdiction", ""),
        ("Signature", ""),
        ("Date", ""),
    ]
    for i, (k, v) in enumerate(rows):
        row = t.rows[i].cells
        row[0].text = ""
        r = row[0].paragraphs[0].add_run(k)
        r.bold = True; r.font.size = Pt(10)
        row[1].text = v


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main() -> None:
    soup = load_source()
    policies = extract_policies(soup)

    doc = Document()
    build_styles(doc)
    _enable_auto_update_fields(doc)

    # Section 1 — cover (no header/footer)
    section = doc.sections[0]
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.different_first_page_header_footer = True
    # First-page header/footer stay empty; body header/footer per below
    _apply_header_footer(section)
    # Blank first-page header/footer
    for target in [section.first_page_header, section.first_page_footer]:
        target.is_linked_to_previous = False
        for p_ in target.paragraphs:
            for r in list(p_.runs):
                r.text = ""

    build_cover(doc)
    _add_page_break(doc)

    build_confidentiality(doc)
    _add_page_break(doc)

    build_executive_summary(doc)
    _add_page_break(doc)

    build_document_control(doc)
    _add_page_break(doc)

    build_marketplace_overview(doc)
    _add_page_break(doc)

    build_legal_review_scope(doc)
    _add_page_break(doc)

    build_launch_readiness_dashboard(doc, policies)
    _add_page_break(doc)

    build_risk_matrix(doc, policies)
    _add_page_break(doc)

    build_counsel_deliverables(doc)
    _add_page_break(doc)

    build_open_legal_questions(doc, policies)
    _add_page_break(doc)

    build_cross_policy_dependency_map(doc)
    _add_page_break(doc)

    build_toc(doc, policies)
    _add_page_break(doc)

    # Policy sections
    for pol in policies:
        build_policy_divider(doc, pol)
        _add_page_break(doc)
        build_policy_overview(doc, pol)
        build_policy_body(doc, pol)
        build_policy_revision_history(doc, pol)
        build_policy_appendix_a(doc, pol)
        _add_page_break(doc)

    build_master_appendix_a(doc, policies)
    _add_page_break(doc)

    build_master_revision_log(doc, policies)
    _add_page_break(doc)

    build_counsel_workbook(doc, policies)
    _add_page_break(doc)

    build_launch_recommendation(doc)
    _add_page_break(doc)

    build_attorney_signoff(doc)

    os.makedirs(os.path.dirname(OUT_FILE), exist_ok=True)
    doc.save(OUT_FILE)
    size = os.path.getsize(OUT_FILE)
    print(f"SUCCESS: {OUT_FILE} — {size} bytes ({size/1024:.1f} KB) · {len(policies)} policies rendered")


if __name__ == "__main__":
    main()
