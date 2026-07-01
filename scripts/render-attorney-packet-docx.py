"""Render the Attorney Review Packet as a DOCX file.

Reads the same HTML extract used by the PDF pipeline
(`/tmp/attorney_packet.json`) and walks the structured
`pkt-*` class markup to build a properly-styled Word
document that counsel can redline.

Env vars:
  IN_FILE  — path to the extracted packet_data.json
             (default /tmp/attorney_packet.json)
  OUT_FILE — output DOCX path (default
             /app/frontend/public/downloads/attorney-review-packet-<DATE>.docx)
"""
from __future__ import annotations

import json
import os
import re
from datetime import date

from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.shared import Inches, Pt, RGBColor

IN_FILE = os.environ.get("IN_FILE", "/tmp/attorney_packet.json")
OUT_FILE = os.environ.get(
    "OUT_FILE",
    f"/app/frontend/public/downloads/attorney-review-packet-{date.today().isoformat()}.docx",
)

# Strip dev-inspector attributes (same as PDF renderer). These are
# HTML attributes only, so bs4 handles them naturally by not
# including attributes in text extraction — but we scrub the raw
# HTML defensively before parsing.
INSPECTOR_ATTR = re.compile(
    r'\s(?:x-source-[a-z-]+|x-file-[a-z-]+|x-line-number|x-column|x-array-var|x-array-index)="[^"]*"'
)


def _text(node) -> str:
    """Collapse whitespace and strip."""
    if node is None:
        return ""
    return re.sub(r"\s+", " ", node.get_text(" ", strip=True)).strip()


def _has_class(tag: Tag, cls: str) -> bool:
    if not isinstance(tag, Tag):
        return False
    classes = tag.get("class", [])
    if isinstance(classes, str):
        classes = classes.split()
    return cls in classes


def _add_styled_paragraph(doc: Document, text: str, *, style: str | None = None,
                          bold: bool = False, italic: bool = False,
                          size: float | None = None, color: RGBColor | None = None,
                          space_after: float | None = None) -> object:
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    if not text:
        return p
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    return p


def _add_page_break(doc: Document) -> None:
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


# ---------------------- Cover Sheet ----------------------


def render_cover(doc: Document, cover: Tag) -> None:
    # Eyebrow
    eyebrow = cover.select_one(".pkt-eyebrow")
    if eyebrow:
        _add_styled_paragraph(
            doc, _text(eyebrow), size=9, bold=True,
            color=RGBColor(0x66, 0x66, 0x66),
        )

    # H1
    h1 = cover.select_one(".pkt-h1")
    if h1:
        p = doc.add_paragraph()
        run = p.add_run(_text(h1))
        run.bold = True
        run.font.size = Pt(22)
        p.paragraph_format.space_after = Pt(4)

    # Meta rows (dl.pkt-meta)
    meta = cover.select_one(".pkt-meta")
    if meta:
        for row in meta.find_all(["dt", "dd"]):
            if row.name == "dt":
                _add_styled_paragraph(
                    doc, _text(row), bold=True, size=9,
                    space_after=0,
                )
            else:
                _add_styled_paragraph(doc, _text(row), size=10, space_after=4)

    # Walk each direct child of the cover section in document order.
    # Skip elements already handled above.
    handled = {eyebrow, h1, meta}
    for child in cover.children:
        if not isinstance(child, Tag):
            continue
        if child in handled:
            continue
        if _has_class(child, "pkt-pagebreak"):
            _add_page_break(doc)
            continue
        if child.name == "h2" or _has_class(child, "pkt-h2"):
            p = doc.add_paragraph()
            run = p.add_run(_text(child))
            run.bold = True
            run.font.size = Pt(14)
            continue
        if child.name == "h3" or _has_class(child, "pkt-h3"):
            p = doc.add_paragraph()
            run = p.add_run(_text(child))
            run.bold = True
            run.font.size = Pt(12)
            continue
        if child.name == "p":
            _add_styled_paragraph(doc, _text(child), size=10)
            continue
        if child.name == "ol":
            for li in child.find_all("li", recursive=False):
                doc.add_paragraph(_text(li), style="List Number")
            continue
        if child.name == "ul":
            for li in child.find_all("li", recursive=False):
                doc.add_paragraph(_text(li), style="List Bullet")
            continue

    _add_page_break(doc)


# ---------------------- Policy Sections ----------------------


def render_policy(doc: Document, policy: Tag) -> None:
    # Policy title
    title = policy.select_one(".pkt-h1-pol") or policy.select_one("h1")
    if title:
        p = doc.add_paragraph()
        run = p.add_run(_text(title))
        run.bold = True
        run.font.size = Pt(18)
        p.paragraph_format.space_after = Pt(6)

    # Header meta (version / effective / last updated)
    meta_inline = policy.select_one(".pkt-meta-inline")
    if meta_inline:
        _add_styled_paragraph(
            doc, _text(meta_inline), size=9, italic=True,
            color=RGBColor(0x55, 0x55, 0x55), space_after=6,
        )

    # Description
    desc = policy.select_one(".pkt-description")
    if desc:
        _add_styled_paragraph(doc, _text(desc), size=10, italic=True, space_after=6)

    # Intro
    intro = policy.select_one(".pkt-intro")
    if intro:
        _add_styled_paragraph(doc, _text(intro), size=10, space_after=8)

    # Table of Contents (small, italic)
    toc = policy.select_one(".pkt-toc")
    if toc:
        _add_styled_paragraph(
            doc, "Table of Contents", bold=True, size=10, space_after=2,
        )
        for li in toc.find_all("li"):
            _add_styled_paragraph(
                doc, "• " + _text(li), size=9,
                color=RGBColor(0x55, 0x55, 0x55), space_after=0,
            )
        _add_styled_paragraph(doc, " ", size=6, space_after=4)

    # Body blocks — each numbered section
    body = policy.select_one(".pkt-body")
    if body:
        for block in body.select(".pkt-body-block"):
            render_body_block(doc, block)

    # Final callout
    callout = policy.select_one(".pkt-callout")
    if callout:
        text = _text(callout)
        if text:
            _add_styled_paragraph(
                doc, text, size=10, italic=True,
                color=RGBColor(0x33, 0x66, 0x99), space_after=8,
            )

    # Version outro
    outro = policy.select_one(".pkt-outro")
    if outro:
        _add_styled_paragraph(
            doc, _text(outro), size=9, italic=True,
            color=RGBColor(0x55, 0x55, 0x55), space_after=8,
        )

    # Revision history
    revision = policy.select_one(".pkt-revision")
    if revision:
        render_meta_block(doc, revision)

    # Related policies
    related = policy.select_one(".pkt-related")
    if related:
        render_meta_block(doc, related)

    # Attorney Review Notes (Appendix A only — in attorney mode B and C are hidden)
    attorney = policy.select_one(".pkt-attorney")
    if attorney:
        render_meta_block(doc, attorney)

    _add_page_break(doc)


def render_block(doc: Document, block: Tag) -> None:
    """Render a single block inside a policy section."""
    if not isinstance(block, Tag):
        return
    if _has_class(block, "pkt-pagebreak"):
        _add_page_break(doc)
        return

    if _has_class(block, "pkt-body-block"):
        render_body_block(doc, block)
        return

    if _has_class(block, "pkt-callout"):
        text = _text(block)
        if text:
            _add_styled_paragraph(
                doc, text, size=10, italic=True,
                color=RGBColor(0x33, 0x66, 0x99), space_after=8,
            )
        return

    if _has_class(block, "pkt-hierarchy") or _has_class(block, "pkt-related") \
       or _has_class(block, "pkt-revision") or _has_class(block, "pkt-attorney"):
        render_meta_block(doc, block)
        return

    # Fallback — headings and paragraphs
    if block.name in ("h1", "h2", "h3", "h4"):
        size_map = {"h1": 16, "h2": 14, "h3": 12, "h4": 11}
        p = doc.add_paragraph()
        run = p.add_run(_text(block))
        run.bold = True
        run.font.size = Pt(size_map[block.name])
        return
    if block.name == "p":
        _add_styled_paragraph(doc, _text(block), size=10, space_after=6)
        return
    if block.name == "ul":
        for li in block.find_all("li", recursive=False):
            doc.add_paragraph(_text(li), style="List Bullet")
        return
    if block.name == "ol":
        for li in block.find_all("li", recursive=False):
            doc.add_paragraph(_text(li), style="List Number")
        return


def render_body_block(doc: Document, block: Tag) -> None:
    """A single §-numbered body block inside a policy."""
    # Heading
    heading = block.select_one(".pkt-h3, .pkt-h4, h3, h4")
    if heading:
        p = doc.add_paragraph()
        run = p.add_run(_text(heading))
        run.bold = True
        run.font.size = Pt(12)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)

    # Direct-child body content
    for el in block.children:
        if not isinstance(el, Tag) or el is heading:
            continue
        if _has_class(el, "pkt-h3") or _has_class(el, "pkt-h4"):
            continue  # already handled
        if el.name == "p" or _has_class(el, "pkt-p"):
            _add_styled_paragraph(doc, _text(el), size=10, space_after=4)
        elif el.name == "ul" or _has_class(el, "pkt-ul"):
            for li in el.find_all("li", recursive=False):
                doc.add_paragraph(_text(li), style="List Bullet")
        elif el.name == "ol" or _has_class(el, "pkt-ol"):
            for li in el.find_all("li", recursive=False):
                doc.add_paragraph(_text(li), style="List Number")
        elif el.name == "dl":
            for row in el.find_all(["dt", "dd"], recursive=False):
                if row.name == "dt":
                    _add_styled_paragraph(
                        doc, _text(row), size=10, bold=True, space_after=0,
                    )
                else:
                    _add_styled_paragraph(
                        doc, _text(row), size=10, space_after=4,
                    )


def render_meta_block(doc: Document, block: Tag) -> None:
    """Hierarchy / Related / Revision / Attorney blocks."""
    heading = block.select_one(".pkt-h3, .pkt-h4, h3, h4")
    if heading:
        p = doc.add_paragraph()
        run = p.add_run(_text(heading))
        run.bold = True
        run.font.size = Pt(11)
        p.paragraph_format.space_before = Pt(4)
    for el in block.children:
        if not isinstance(el, Tag) or el is heading:
            continue
        if _has_class(el, "pkt-h3") or _has_class(el, "pkt-h4"):
            continue
        if el.name == "p" or _has_class(el, "pkt-p"):
            _add_styled_paragraph(doc, _text(el), size=10, space_after=3)
        elif el.name == "ul" or _has_class(el, "pkt-ul"):
            for li in el.find_all("li"):
                doc.add_paragraph(_text(li), style="List Bullet")
        elif el.name == "ol" or _has_class(el, "pkt-ol"):
            for li in el.find_all("li"):
                doc.add_paragraph(_text(li), style="List Number")
        elif el.name == "h4" or _has_class(el, "pkt-h4"):
            # sub-heading (e.g., "Appendix A — Attorney Review Notes")
            _add_styled_paragraph(
                doc, _text(el), bold=True, size=10, space_after=2,
            )


# ---------------------- Glossary ----------------------


def render_glossary(doc: Document, glossary: Tag) -> None:
    heading = glossary.find(["h1", "h2"])
    if heading:
        p = doc.add_paragraph()
        run = p.add_run(_text(heading))
        run.bold = True
        run.font.size = Pt(16)
        p.paragraph_format.space_after = Pt(6)
    for row in glossary.find_all("dt"):
        term = _text(row)
        defn_tag = row.find_next_sibling("dd")
        defn = _text(defn_tag) if defn_tag else ""
        p = doc.add_paragraph()
        r1 = p.add_run(f"{term} — ")
        r1.bold = True
        r1.font.size = Pt(10)
        r2 = p.add_run(defn)
        r2.font.size = Pt(10)


# ---------------------- Attorney Response Sheet (prelude) ----------------------


def render_response_sheet(doc: Document) -> None:
    """One-page fill-in sheet for counsel to return with their notes."""
    p = doc.add_paragraph()
    run = p.add_run("Attorney Response Sheet")
    run.bold = True
    run.font.size = Pt(20)
    p.paragraph_format.space_after = Pt(6)

    _add_styled_paragraph(
        doc,
        "Please complete the table below and return with your notes. Redline the "
        "body of this document directly; use this sheet as a summary layer for "
        "our operations team to triage the response.",
        size=10, italic=True, space_after=10,
    )

    docs = [
        ("Terms of Service", "2.6"),
        ("Maker Agreement", "3.6"),
        ("Privacy Policy", "3.4"),
        ("Cookies Policy", "1.1"),
        ("Buyer Protection Policy", "1.1"),
        ("Returns & Refunds Policy", "3.4"),
        ("Shipping & Logistics Policy", "3.1"),
        ("Prohibited Items Policy", "3.3"),
        ("Community Guidelines", "1.0"),
        ("Intellectual Property & DMCA Policy", "3.1"),
        ("Fee & Pricing Policy", "1.3"),
        ("Marketplace Promise", "1.0"),
        ("Accessibility Statement", "1.0"),
        ("Creator-Owned AI Policy", "1.0"),
    ]

    table = doc.add_table(rows=1 + len(docs), cols=5)
    table.style = "Light Grid"
    hdr = table.rows[0].cells
    for i, label in enumerate(["Document", "Version", "Sign-off (Y/N)",
                               "Blockers", "Est. hours"]):
        cell = hdr[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(label)
        run.bold = True
        run.font.size = Pt(9)

    for i, (name, ver) in enumerate(docs, start=1):
        row = table.rows[i].cells
        row[0].text = name
        row[1].text = ver
        # leave 2, 3, 4 blank for counsel
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for r in para.runs:
                    if not r.bold:
                        r.font.size = Pt(9)

    _add_styled_paragraph(
        doc, " ", size=8, space_after=4,
    )
    _add_styled_paragraph(
        doc,
        "General notes (blockers, cross-doc concerns, additional recommendations):",
        size=10, bold=True, space_after=2,
    )
    for _ in range(6):
        _add_styled_paragraph(
            doc,
            "____________________________________________________________________________________________",
            size=10, space_after=4,
        )

    _add_page_break(doc)


# ---------------------- Main ----------------------


def main() -> None:
    with open(IN_FILE) as f:
        data = json.load(f)

    html = INSPECTOR_ATTR.sub("", data["html"])
    soup = BeautifulSoup(html, "lxml")

    doc = Document()

    # Standardize document defaults
    style = doc.styles["Normal"]
    style.font.name = "Georgia"
    style.font.size = Pt(10)

    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

    # Attorney Response Sheet at the top (attorney-only DOCX)
    render_response_sheet(doc)

    # Cover sheet
    cover = soup.select_one(".pkt-cover")
    if cover:
        render_cover(doc, cover)

    # Every policy section
    for policy_section in soup.select(".pkt-policy"):
        render_policy(doc, policy_section)

    # Glossary appendix
    glossary = soup.select_one(".pkt-glossary")
    if glossary:
        render_glossary(doc, glossary)

    # Footer paragraph
    footer = soup.select_one(".pkt-footer")
    if footer:
        _add_styled_paragraph(
            doc, _text(footer), size=8, italic=True,
            color=RGBColor(0x66, 0x66, 0x66),
            space_after=0,
        )

    os.makedirs(os.path.dirname(OUT_FILE), exist_ok=True)
    doc.save(OUT_FILE)
    size = os.path.getsize(OUT_FILE)
    print(f"SUCCESS: {OUT_FILE} — {size} bytes ({size/1024:.1f} KB)")


if __name__ == "__main__":
    main()
